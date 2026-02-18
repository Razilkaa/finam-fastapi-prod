"""Generate quotes Word document by filling a .docx table template."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.shared import RGBColor


GREEN = RGBColor.from_string("00B050")
RED = RGBColor.from_string("FF0000")


# Labels must match the first column text in Template_quotes.docx.
SYMBOL_TO_TEMPLATE_LABEL: dict[str, str] = {
    # FX
    "dxy": "Индекс USD",
    "eurusd": "EUR/USD",
    "gbpusd": "GBP/USD",
    "usdjpy": "USD/JPY",
    "usdrub": "USD/RUB",
    "usdcny": "USD/CNY",
    "usdinr": "USD/INR",
    "usdbrl": "USD/BRL",
    # Commodities
    "brent": "Фьючерс на нефть Brent",
    "crude oil": "Фьючерсы на нефть WTI",
    "gold": "Фьючерс на золото",
    "silver": "Фьючерс на серебро",
    "copper": "Фьючерс на медь",
    "nickel": "Фьючерс на никель",
    "aluminum": "Алюминий",
}


@dataclass(frozen=True)
class Quote:
    symbol: str
    old_price: Optional[float]
    new_price_raw: Optional[str]
    pct_change: Optional[float]
    report_date: Optional[str]


def _to_float(value: Any) -> Optional[float]:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        s = value.strip().replace("\u00A0", " ")
        s = s.replace("%", "")
        s = s.replace(" ", "")
        s = s.replace(",", ".")
        if not s:
            return None
        try:
            return float(s)
        except ValueError:
            return None
    return None


def _format_price(value: Optional[str]) -> str:
    if value is None:
        return ""
    s = str(value).strip()
    if not s:
        return ""
    return s.replace(".", ",")


def _format_pct(value: Optional[float]) -> str:
    if value is None:
        return ""
    return f"{value:.2f}%".replace(".", ",")


def _pct_color(value: Optional[float]) -> Optional[RGBColor]:
    if value is None:
        return None
    if value > 0:
        return GREEN
    if value < 0:
        return RED
    return None


def _set_cell_text(cell, text: str, *, color: Optional[RGBColor] = None) -> None:
    paragraphs = cell.paragraphs
    if not paragraphs:
        p = cell.add_paragraph()
    else:
        p = paragraphs[0]

    run = None
    for r in p.runs:
        if run is None:
            run = r
        else:
            r.text = ""

    if run is None:
        run = p.add_run()

    run.text = text
    if color is not None:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = None


def _parse_report_date(value: Any) -> Optional[date]:
    if value is None:
        return None
    if isinstance(value, (datetime, date)):
        return value.date() if isinstance(value, datetime) else value

    s = str(value).strip()
    if not s:
        return None

    if s.endswith("Z"):
        try:
            return datetime.fromisoformat(s.replace("Z", "+00:00")).date()
        except ValueError:
            pass

    try:
        return datetime.fromisoformat(s).date()
    except ValueError:
        pass

    try:
        return datetime.strptime(s, "%d.%m.%Y").date()
    except ValueError:
        return None


def parse_quotes(payload: list[dict]) -> tuple[list[Quote], Optional[date]]:
    quotes: list[Quote] = []
    report_dt: Optional[date] = None

    for item in payload:
        if not isinstance(item, dict):
            continue

        report_date_raw = item.get("report_date")
        if report_dt is None:
            report_dt = _parse_report_date(report_date_raw)

        symbol = str(item.get("symbol", "")).strip()
        if not symbol:
            continue

        old_price = _to_float(item.get("old_price"))
        new_price_raw = item.get("new_price")
        new_price_str = None if new_price_raw is None else str(new_price_raw).strip()
        pct = _to_float(item.get("pct_change"))

        if pct is None and old_price is not None:
            new_price_num = _to_float(new_price_str)
            if new_price_num is not None and old_price != 0:
                pct = (new_price_num - old_price) / old_price * 100.0

        quotes.append(
            Quote(
                symbol=symbol,
                old_price=old_price,
                new_price_raw=new_price_str,
                pct_change=pct,
                report_date=None if report_date_raw is None else str(report_date_raw).strip(),
            )
        )

    return quotes, report_dt


def _build_label_index(table) -> dict[str, int]:
    index: dict[str, int] = {}
    for ri, row in enumerate(table.rows):
        label = row.cells[0].text.strip()
        if not label:
            continue
        index[label] = ri
    return index


def fill_template(*, template_path: Path, quotes: list[Quote]) -> tuple[BytesIO, int]:
    quotes_by_symbol = {q.symbol.strip().lower(): q for q in quotes}

    doc = Document(str(template_path))
    if not doc.tables:
        raise ValueError("Template must contain at least one table")

    table = doc.tables[0]
    label_to_row = _build_label_index(table)

    updated = 0
    for symbol_key, label in SYMBOL_TO_TEMPLATE_LABEL.items():
        quote = quotes_by_symbol.get(symbol_key)
        if quote is None:
            continue
        row_idx = label_to_row.get(label)
        if row_idx is None:
            continue

        price = _format_price(quote.new_price_raw)
        pct_text = _format_pct(quote.pct_change)
        color = _pct_color(quote.pct_change)

        _set_cell_text(table.cell(row_idx, 1), price)
        _set_cell_text(table.cell(row_idx, 2), pct_text, color=color)
        updated += 1

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, updated


def get_quotes_filename(report_dt: Optional[date]) -> str:
    if report_dt is not None:
        return f"Daily_quotes_{report_dt.strftime('%d.%m.%Y')}.docx"
    return "Daily_quotes.docx"
