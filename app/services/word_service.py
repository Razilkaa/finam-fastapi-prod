"""Word document generation service."""
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Optional, Iterator

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

from app.utils.date_utils import (
    get_monday_of_week,
    get_week_dates,
    choose_reference_monday,
    group_items_by_date,
    parse_time_for_sort,
)
from app.utils.constants import COUNTRY_NAMES_RU, COUNTRY_NAMES_EN


# Дни недели
DAYS_RU = ["Понедельник", "Вторник", "Среда", "Четверг", "Пятница", "Суббота", "Воскресенье"]
DAYS_EN = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]

# Месяцы
MONTHS_RU = ["января", "февраля", "марта", "апреля", "мая", "июня",
             "июля", "августа", "сентября", "октября", "ноября", "декабря"]
MONTHS_EN = ["January", "February", "March", "April", "May", "June",
             "July", "August", "September", "October", "November", "December"]


def convert_to_24h(time_str: str) -> str:
    """Конвертация времени из AM/PM в 24-часовой формат."""
    if not time_str or not isinstance(time_str, str):
        return ""
    
    time_str = time_str.strip().upper()
    if not time_str:
        return ""
    
    if "AM" not in time_str and "PM" not in time_str:
        return time_str
    
    is_pm = "PM" in time_str
    time_clean = time_str.replace("AM", "").replace("PM", "").strip()
    
    try:
        if ":" in time_clean:
            parts = time_clean.split(":")
            hours = int(parts[0])
            minutes = int(parts[1]) if len(parts) > 1 else 0
            
            if is_pm and hours != 12:
                hours += 12
            elif not is_pm and hours == 12:
                hours = 0
            
            return f"{hours:02d}:{minutes:02d}"
    except (ValueError, IndexError):
        pass
    
    return time_str


def format_date_header(d: date, lang: str) -> str:
    """Форматирование заголовка дня: 'Понедельник, 12 января'."""
    if lang == "ru":
        return f"{DAYS_RU[d.weekday()]}, {d.day} {MONTHS_RU[d.month - 1]}"
    else:
        return f"{DAYS_EN[d.weekday()]}, {MONTHS_EN[d.month - 1]} {d.day}"


def format_event_line(time_str: str, country: str, event: str, lang: str) -> str:
    """Форматирование строки события: '02:50 – США: событие'."""
    country_names = COUNTRY_NAMES_RU if lang == "ru" else COUNTRY_NAMES_EN
    country_display = country_names.get(country, country)
    
    time_24 = convert_to_24h(time_str)
    
    if time_24:
        return f"{time_24} – {country_display}: {event}"
    else:
        return f"{country_display}: {event}"


def format_holiday_line(holidays: list[dict], lang: str) -> str:
    """Форматирование строки праздников."""
    country_names = COUNTRY_NAMES_RU if lang == "ru" else COUNTRY_NAMES_EN
    
    grouped: dict[str, list[str]] = {}
    for hol in holidays:
        name = hol.get("holiday", "") or hol.get("event", "")
        country = hol.get("country", "")
        if name:
            grouped.setdefault(name, []).append(country)
    
    parts = []
    for name, countries in grouped.items():
        country_list = ", ".join(country_names.get(c, c) for c in sorted(set(countries)))
        if lang == "ru":
            parts.append(f"{name}. Праздник в {country_list}")
        else:
            parts.append(f"{name}. Markets in {country_list}")
    
    return "; ".join(parts)


def generate_content(
    events: list[dict],
    holidays: list[dict],
    lang: str,
    monday: Optional[date] = None,
) -> str:
    """Генерация текстового контента для одного языка."""
    
    events_by_date = group_items_by_date(events)
    holidays_by_date = group_items_by_date(holidays)

    if monday is None:
        monday = choose_reference_monday(events_by_date, holidays_by_date)
    
    week_dates = get_week_dates(monday)
    
    lines = []
    no_data_text = "Нет важных макроданных" if lang == "ru" else "No important macroeconomic data"
    
    for d in week_dates:
        lines.append(format_date_header(d, lang))
        
        day_holidays = holidays_by_date.get(d, [])
        day_events = events_by_date.get(d, [])
        
        has_content = False
        
        if day_holidays:
            lines.append(format_holiday_line(day_holidays, lang))
            has_content = True
        
        if day_events:
            day_events.sort(key=lambda x: parse_time_for_sort(x.get("time", "")))
            for ev in day_events:
                line = format_event_line(
                    ev.get("time", ""),
                    ev.get("country", ""),
                    ev.get("event", ""),
                    lang
                )
                lines.append(line)
            has_content = True
        
        if not has_content:
            lines.append(no_data_text)
        
        lines.append("")
    
    return "\n".join(lines).strip()


def insert_paragraph_after(paragraph: Paragraph, parent) -> Paragraph:
    """Вставить новый параграф ПОСЛЕ указанного."""
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, parent)


def iter_all_paragraphs(doc: Document) -> Iterator[tuple[Paragraph, any]]:
    """Итерация по ВСЕМ параграфам документа."""
    for para in doc.paragraphs:
        yield para, doc
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para, cell
    
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                yield para, section.header
        if section.footer:
            for para in section.footer.paragraphs:
                yield para, section.footer


def format_paragraph(para: Paragraph, line: str, font_name: str, font_size):
    """Форматирование параграфа с текстом."""
    is_day_header = any(day in line for day in DAYS_RU + DAYS_EN)
    
    if is_day_header:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(11)
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(line)
        run.bold = False
        run.font.name = font_name or "Arial"
        run.font.size = font_size or Pt(11)


def replace_placeholder(doc: Document, placeholder: str, content: str) -> bool:
    """Замена placeholder-а в документе на контент."""
    for para, parent in iter_all_paragraphs(doc):
        if placeholder in para.text:
            if para.runs:
                first_run = para.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
            else:
                font_name = "Arial"
                font_size = Pt(11)
            
            para.clear()
            
            content_lines = content.split("\n")
            
            if content_lines:
                format_paragraph(para, content_lines[0], font_name, font_size)
            
            current_para = para
            for line in content_lines[1:]:
                new_para = insert_paragraph_after(current_para, parent)
                format_paragraph(new_para, line, font_name, font_size)
                current_para = new_para
            
            return True
    
    return False


def replace_inline_placeholder(doc: Document, placeholder: str, value: str) -> bool:
    """Замена placeholder-а на значение внутри существующего абзаца."""
    replaced = False
    for para, _parent in iter_all_paragraphs(doc):
        if placeholder in para.text:
            new_text = para.text.replace(placeholder, value)
            if para.runs:
                first_run = para.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
            else:
                font_name = "Arial"
                font_size = Pt(11)

            para.clear()
            run = para.add_run(new_text)
            run.font.name = font_name or "Arial"
            run.font.size = font_size or Pt(11)
            replaced = True
    return replaced


def generate_word(
    work_en: list[dict],
    work_ru: list[dict],
    holidays_en: list[dict],
    holidays_ru: list[dict],
    template_path: str | Path
) -> BytesIO:
    """Генерация Word документа из шаблона."""
    template_path = Path(template_path)
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")
    
    doc = Document(str(template_path))
    
    combined_events_by_date = group_items_by_date(work_en + work_ru)
    combined_holidays_by_date = group_items_by_date(holidays_en + holidays_ru)
    monday = choose_reference_monday(combined_events_by_date, combined_holidays_by_date)
    calendar_date = f"{monday.day:02d}.{monday.month:02d}.{str(monday.year)[2:]}"

    content_ru = generate_content(work_ru, holidays_ru, "ru", monday=monday)
    content_en = generate_content(work_en, holidays_en, "en", monday=monday)
    
    if not replace_placeholder(doc, "{{CONTENT_RU}}", content_ru):
        raise ValueError("Placeholder {{CONTENT_RU}} not found in template.")
    
    if not replace_placeholder(doc, "{{CONTENT_EN}}", content_en):
        raise ValueError("Placeholder {{CONTENT_EN}} not found in template.")

    replace_inline_placeholder(doc, "{{CALENDAR_DATE}}", calendar_date)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def get_output_filename(events: list[dict], holidays: Optional[list[dict]] = None) -> str:
    """Генерация имени выходного файла на основе данных."""
    events_by_date = group_items_by_date(events)
    holidays_by_date = group_items_by_date(holidays or [])
    monday = choose_reference_monday(events_by_date, holidays_by_date)
    return f"Calendar_{monday.day:02d}.{monday.month:02d}.{monday.year}.docx"
